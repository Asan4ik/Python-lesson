from openpyxl import Workbook
import json
import requests
import docx

# 1
wb = Workbook()
ws = wb.active

letter_str = ".abcdefghijklmnop"

# а)
for i in range(1, 11):
    ws[f"{letter_str[i]}1"] = i
    ws[f"{letter_str[i]}2"] = letter_str[i]

# б)
wb.save("test.xlsx")

wb2 = Workbook()
ws2 = wb2.active

# в)
for i in range(1, 11):
    ws2[f"A{i}"] = i
    ws2[f"B{i}"] = letter_str[i]

wb2.save("test2.xlsx")


# а)
response = requests.get("https://jsonplaceholder.typicode.com/todos/1")
data = response.json()

with open("data.json", "w") as file:
    json.dump(data, file)

# б)
with open("data.json", "r") as file:
    my_dict = json.load(file)

print(my_dict)

# в)
dict_list = [my_dict.copy() for i in range(100)]

with open("a_lot_of_data.json", "w") as file:
    json.dump(dict_list, file)



# а)
doc = docx.Document()
doc.add_paragraph("Hello python")
doc.save("hello_python.docx")
# б)
print(doc.paragraphs[0].text)

# в)
doc2 = docx.Document()
doc2.add_paragraph("Paragraph 1")
doc2.add_paragraph("Paragraph 2")
doc2.save("paragraphs.docx")